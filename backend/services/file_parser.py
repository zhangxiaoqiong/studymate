"""
文件解析服务

支持多种文件格式的文本提取：
- .txt：纯文本文件（支持 UTF-8 和 GBK 编码）
- .md：Markdown 文件
- .pdf：PDF 文件（使用 pdfplumber）
- .docx/.doc：Word 文档（使用 python-docx）

限制：
- 最大文件大小：5MB
- 支持的格式：txt, pdf, docx, doc, md
"""
import pdfplumber
from docx import Document
from pathlib import Path
from typing import Tuple

class FileParser:
    """文件解析器，支持 .txt, .pdf, .docx, .md 等多种格式"""

    MAX_FILE_SIZE = 5 * 1024 * 1024  # 5MB
    SUPPORTED_FORMATS = {'.txt', '.pdf', '.docx', '.doc', '.md'}

    @staticmethod
    def validate_file(file_path: str, file_size: int) -> Tuple[bool, str]:
        """
        验证文件格式和大小

        Args:
            file_path: 文件路径
            file_size: 文件大小（字节）

        Returns:
            tuple: (是否有效, 错误信息)
        """
        path = Path(file_path)
        suffix = path.suffix.lower()

        if suffix not in FileParser.SUPPORTED_FORMATS:
            return False, f"不支持的文件格式: {suffix}。支持的格式: {', '.join(FileParser.SUPPORTED_FORMATS)}"

        if file_size > FileParser.MAX_FILE_SIZE:
            return False, f"文件过大。最大支持 {FileParser.MAX_FILE_SIZE / (1024*1024):.0f}MB"

        return True, ""

    @staticmethod
    def parse_txt(file_path: str) -> str:
        """
        解析纯文本文件，自动尝试 UTF-8 和 GBK 编码

        Args:
            file_path: 文件路径

        Returns:
            str: 文件内容
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            return content.strip()
        except UnicodeDecodeError:
            # 尝试用其他编码
            with open(file_path, 'r', encoding='gbk') as f:
                content = f.read()
            return content.strip()

    @staticmethod
    def parse_pdf(file_path: str) -> str:
        """
        解析 PDF 文件，提取所有页面的文本

        Args:
            file_path: 文件路径

        Returns:
            str: 合并后的文本内容

        Raises:
            ValueError: 如果 PDF 解析失败
        """
        text = []
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text.append(page_text)
            return '\n'.join(text).strip()
        except Exception as e:
            raise ValueError(f"PDF 解析失败: {str(e)}")

    @staticmethod
    def parse_docx(file_path: str) -> str:
        """
        解析 DOCX 文件，提取段落和表格内容

        Args:
            file_path: 文件路径

        Returns:
            str: 提取的文本内容

        Raises:
            ValueError: 如果 DOCX 解析失败
        """
        text = []
        try:
            doc = Document(file_path)
            # 提取段落
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text.append(paragraph.text)

            # 提取表格
            for table in doc.tables:
                for row in table.rows:
                    row_text = '\t'.join(cell.text for cell in row.cells)
                    if row_text.strip():
                        text.append(row_text)

            return '\n'.join(text).strip()
        except Exception as e:
            raise ValueError(f"DOCX 解析失败: {str(e)}")

    @staticmethod
    def parse_file(file_path: str) -> str:
        """
        根据文件类型自动选择解析器

        Args:
            file_path: 文件路径

        Returns:
            str: 提取的文本内容

        Raises:
            ValueError: 如果文件格式不支持或解析失败
        """
        path = Path(file_path)
        suffix = path.suffix.lower()

        try:
            if suffix in {'.txt', '.md'}:
                return FileParser.parse_txt(file_path)
            elif suffix == '.pdf':
                return FileParser.parse_pdf(file_path)
            elif suffix in {'.docx', '.doc'}:
                return FileParser.parse_docx(file_path)
            else:
                raise ValueError(f"不支持的文件格式: {suffix}")
        except ValueError:
            raise
        except Exception as e:
            raise ValueError(f"文件解析失败: {str(e)}")
